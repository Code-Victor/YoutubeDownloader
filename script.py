from pytube import YouTube,Playlist,extract
from youtube_transcript_api import YouTubeTranscriptApi as ytTrans
from youtube_transcript_api.formatters import WebVTTFormatter
from pytube.cli import on_progress
from docx import Document
import os


class GetYoutubeVideo:
    def __init__(self,url,destinaiton=r"C:\users\oluwa\tutorials",audio=False,resolution='720p',Transcript=True,language=['en']):
        self.url=url
        self.destination=destinaiton
        self.transcript=Transcript
        self.resolution=resolution
        self.audio=audio
        self.transcript_language=language
        self.videoId=extract.video_id(self.url)
        self.yt=YouTube(self.url,on_progress_callback=on_progress)
        self.name=self.yt.title
        for word, initial in {'?':'.','<':'.','>':'.','|':'.','*':'.'}.items():
            self.name = self.name.replace(word.lower(), initial)
        




    def get_transcript(self):
        print("downloading {}.vtt".format(self.name))
        transcript=ytTrans.get_transcript(self.videoId,languages=self.transcript_language)
        webvtt_formatter=WebVTTFormatter()
        webvtt_formatted=webvtt_formatter.format_transcript(transcript=transcript)
        with open('./test/{}.vtt'.format(self.name), 'w', encoding='utf-8') as webvtt_file:
            webvtt_file.write(webvtt_formatted)
        print("downloaded {}.vtt".format(self.name))

    def _get_description(self):

        print("downloading {}.docx".format(self.name))
        document=Document()
        description=self.yt.description
        document.add_heading(self.name,0)
        document.add_paragraph(description)
        document.save('./test/{}.docx'.format(self.name))
        stream=self.yt.streams.filter(res=self.resolution)
        print("downloaded {}.docx".format(self.name))
         
        

    def get_video(self):
        stream=self.yt.streams.filter(res=self.resolution,progressive=True)  
        stream.first().download('./test/')
        print("downloaded {}".format(self.name))


    def all_together(self):
        self.get_transcript()
        self._get_description()
        self.get_video()
    

# youtube=getyoutubevideo(url="https://www.youtube.com/watch?v=VEQaH4LruUo")
# youtube.alltogether()

# print(YouTube("https://youtu.be/_YPScrckx28").streams.filter(progressive=True).get_highest_resolution().resolution)
youtube=GetYoutubeVideo(url="https://www.youtube.com/watch?v=zrs7u6bdbUw")
youtube.all_together()
